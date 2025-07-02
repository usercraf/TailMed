from unicodedata import category

from docx import Document
from docx.shared import Inches
import os


def generate_tailmed_docx(data: dict, output_path: str, doc_name, details, category, pets_name, species, breed, owner_name, owner_phone, create_at):
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()
    # Додаємо картинку зверху
    doc.add_picture("logo.jpg", width=Inches(6.0))
    doc.add_heading("TailMed — Медична форма", 0)
    doc.add_paragraph(f"Ім’я тварини: {pets_name.capitalize()}")
    doc.add_paragraph(f"Вид: {species.capitalize()}")
    doc.add_paragraph(f"Порода: {breed.capitalize()}")
    doc.add_paragraph(f"Власник: {owner_name.capitalize()} ({owner_phone})")
    doc.add_paragraph(f"Дата: {create_at}")

    if category in ['огляд', 'рекомендація']:
        doc.add_heading(f"{data.get('title').capitalize()}:", level=1)
        doc.add_paragraph(details)

    # if data.get("tests"):
    #     doc.add_heading("Направлення на аналізи", level=1)
    #     for test in data["tests"]:
    #         doc.add_paragraph(f"- {test}", style="List Bullet")
    #
    # if data.get("medications"):
    #     doc.add_heading("Рецепт", level=1)
    #     table = doc.add_table(rows=1, cols=3)
    #     table.style = 'Table Grid'
    #     hdr_cells = table.rows[0].cells
    #     hdr_cells[0].text = 'Препарат'
    #     hdr_cells[1].text = 'Дозування'
    #     hdr_cells[2].text = 'Курс'
    #
    #     for med in data["medications"]:
    #         row = table.add_row().cells
    #         row[0].text = med["name"]
    #         row[1].text = med["dose"]
    #         row[2].text = med["duration"]

    doc.add_paragraph(f"\nЛікар: {doc_name}")
    doc.add_paragraph("Підпис: ____________________________")

    doc.save(output_path)